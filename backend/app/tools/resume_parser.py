"""
Resume Text Extraction — Converting PDF/DOCX files to raw text.

This is a TOOL, not an agent. It does one simple thing: take a file,
return its text content. No LLM involved — just parsing.

Why pdfplumber over PyPDF2?
    - Better at handling multi-column layouts (common in resumes)
    - Preserves table structure when extracting
    - More accurate text positioning
    - Handles unicode and special characters better

Why python-docx?
    - The standard library for reading .docx files
    - Handles paragraphs, tables, and headers properly
    - Much more reliable than trying to convert DOCX to PDF first
"""

from pathlib import Path

import pdfplumber
from docx import Document


def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text content from a PDF file.
    
    Goes through each page, extracts text, and joins them with
    double newlines to preserve page separation.
    
    Args:
        file_path: Absolute or relative path to the PDF file
        
    Returns:
        The full text content of the PDF
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        pdfplumber.exceptions.PDFSyntaxError: If the file is corrupted
    """
    text_parts = []
    
    with pdfplumber.open(file_path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text()
            
            if page_text and page_text.strip():
                text_parts.append(page_text.strip())
    
    if not text_parts:
        raise ValueError(
            f"Could not extract any text from '{file_path}'. "
            "The PDF might be image-based (scanned). "
            "Try uploading a text-based PDF or a DOCX file."
        )
    
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text content from a DOCX file.
    
    Reads paragraphs and table cells, joining everything into
    a clean text representation.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        The full text content of the document
    """
    doc = Document(file_path)
    text_parts = []
    
    # Extract text from paragraphs (the main content)
    for paragraph in doc.paragraphs:
        stripped = paragraph.text.strip()
        if stripped:
            text_parts.append(stripped)
    
    # Also extract text from tables — resumes sometimes use tables
    # for layout (e.g., skills in a two-column table)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    if not text_parts:
        raise ValueError(
            f"Could not extract any text from '{file_path}'. "
            "The document appears to be empty."
        )
    
    return "\n".join(text_parts)


def extract_resume_text(file_path: str) -> str:
    """Extract text from a resume file — supports PDF and DOCX.
    
    This is the main entry point for resume text extraction.
    It detects the file format and delegates to the appropriate parser.
    
    Args:
        file_path: Path to the resume file (PDF or DOCX)
        
    Returns:
        Raw text content of the resume
        
    Raises:
        ValueError: If the file format is not supported
        FileNotFoundError: If the file doesn't exist
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {file_path}")
    
    extension = path.suffix.lower()
    
    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif extension in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif extension == ".txt":
        # Fallback: some people might paste text into a .txt file
        return path.read_text(encoding="utf-8")
    else:
        raise ValueError(
            f"Unsupported file format: '{extension}'. "
            "Please upload a PDF, DOCX, or TXT file."
        )
